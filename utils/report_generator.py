from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import datetime
from utils.logger import logger

def set_cell_margins(cell, **kwargs):
    """
    Sets cell margins for a docx table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(kwargs[m]))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def generate_docx_report(image_buffer, prediction_label, confidence_score, confidence_percent, gradcam_buffer=None):
    """
    Generates a professional DOCX report for thyroid cancer detection.
    """
    logger.info(f"Generating enhanced medical report for: {prediction_label}")
    doc = Document()
    
    # 1. Header Branded Bar
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run("THYROCHECK AI | CLINICAL DIAGNOSTICS")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 192, 163) # Emerald
    
    # 2. Title & Date
    title = doc.add_heading('Diagnostic Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp = datetime.datetime.now().strftime('%B %d, %Y | %H:%M:%S')
    run = meta_para.add_run(f"Report Generated: {timestamp}")
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3. Patient Scan & Summary Section
    doc.add_heading('1. Analysis Executive Summary', level=1)
    
    # Summary Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    def fill_row(row_idx, label, value, is_bold_val=False):
        row = table.rows[row_idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = str(value)
        if is_bold_val:
            row.cells[1].paragraphs[0].runs[0].font.bold = True
        set_cell_margins(row.cells[0], start=100)
        set_cell_margins(row.cells[1], start=100)

    fill_row(0, "Diagnostic Determination", prediction_label.upper(), True)
    fill_row(1, "Probability Coefficient", f"{confidence_score:.4f}")
    fill_row(2, "Confidence Level", f"{confidence_percent:.2f}%")
    fill_row(3, "Neural Network ID", "FibonacciNet-v1 (ResNet-DWSC)")

    # Color the determination cell
    if "Malignant" in prediction_label:
        table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 82, 82)
    else:
        table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 192, 163)

    doc.add_paragraph() # Spacer
    
    # 4. Medical Imaging Section
    doc.add_heading('2. Diagnostic Imaging & Interpretability', level=1)
    
    # Create side-by-side or stacked layout
    # Since side-by-side can be tricky in docx without complex tables, we'll do structured stacking
    
    if image_buffer:
        doc.add_heading('Primary Ultrasound / Pathology Scan', level=2)
        image_buffer.seek(0)
        doc.add_picture(image_buffer, width=Inches(4.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    if gradcam_buffer:
        doc.add_heading('Grad-CAM Attention Heatmap', level=2)
        doc.add_paragraph(
            "The following visualization highlights the localized regions (indicated by warmer colors) "
            "that contributed most significantly to the model's classification output."
        )
        gradcam_buffer.seek(0)
        doc.add_picture(gradcam_buffer, width=Inches(4.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. Technical Synopsis
    doc.add_page_break()
    doc.add_heading('3. Technical Methodology', level=1)
    doc.add_paragraph(
        "This diagnostic analysis was performed using ThyroCheck AI's proprietary FibonacciNet architecture. "
        "The model utilizes Fibonacci-scaled filter counts (21, 34, 55, 89, 144, 233, 377) and Partial Connection "
        "Blocks (PCB) for optimal feature extraction from medical ultrasound signals. Interpretability is "
        "provided via Gradient-weighted Class Activation Mapping (Grad-CAM)."
    )
    
    # 6. Disclaimer Header
    doc.add_paragraph()
    doc.add_paragraph("-" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run("IMPORTANT CLINICAL DISCLAIMER")
    run.bold = True
    run.font.size = Pt(11)
    
    p = doc.add_paragraph(
        "This report is generated by an artificial intelligence system and intended for assistive decision-making "
        "only. It does not constitute a final medical diagnosis. All findings should be reviewed and verified by a "
        "qualified radiologist or medical professional."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(9)
    p.style.font.italic = True

    # Save to buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer
